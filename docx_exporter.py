from docx import Document
from docx.shared import Inches
from PyQt6.QtWidgets import QLabel, QTextEdit
from PyQt6.QtCore import Qt

class DocxExporter:
    def __init__(self):
        self.document = Document()
        
    def export(self, canvas, file_path):
        # Create a new document
        doc = Document()
        
        # Iterate through canvas items
        for item in canvas.items:
            if isinstance(item, QLabel):  # Screenshot
                image = item.pixmap()
                # Save image temporarily and add to document
                temp_path = "temp_image.png"
                image.save(temp_path)
                doc.add_picture(temp_path, width=Inches(6))
                import os
                os.remove(temp_path)
                
            elif isinstance(item, QTextEdit):  # Text or code
                doc.add_paragraph(item.toPlainText())
            
            # Add spacing between elements
            doc.add_paragraph()
            
        # Save the document
        doc.save(file_path)